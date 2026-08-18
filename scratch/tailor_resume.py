import docx
import sys

def create_tailored_resume(template_path, output_path):
    doc = docx.Document(template_path)
    
    doc.paragraphs[0].clear().add_run("Haile SOTOME").bold = True
    doc.paragraphs[1].clear().add_run("Montreuil, France")
    doc.paragraphs[2].clear().add_run("+33 6 95 59 86 39 | haile.sotome@skema.edu")
    
    p4 = doc.paragraphs[4]
    p4.clear()
    p4.add_run("SKEMA Business School - Soochow University\tSuzhou, China / Paris, France").bold = True
    
    doc.paragraphs[5].clear().add_run("Masters PGE: Specialization in Corporate Financial Management (CFM)\tExpected 2027")
    
    p6 = doc.paragraphs[6]
    p6.clear()
    p6.add_run("SKEMA Business School - North Carolina State University\tLille, France / Raleigh, US").bold = True
    
    doc.paragraphs[7].clear().add_run("Masters PGE: Specialization in Finance & Quants\t2023 - 2026")
    
    p8 = doc.paragraphs[8]
    p8.clear()
    p8.add_run("Relevant Coursework: ").bold = True
    p8.add_run("Corporate Valuation, Financial Modelling, Cash Flow Forecasting, Risk Management, Big Data Analysis, Derivatives, Structured Products")
    
    p10 = doc.paragraphs[10]
    p10.clear()
    p10.add_run("Université Paris-Est Créteil (UPEC)\tCréteil, France").bold = True
    
    doc.paragraphs[11].clear().add_run("Financial Analyst Intern, Financial Affairs Division\tMay 2025 - Dec 2025")
    doc.paragraphs[12].clear().add_run("Monitored a €330M annual budget and streamlined financial reporting processes.")
    doc.paragraphs[13].clear()
    doc.paragraphs[14].clear()
    
    doc.paragraphs[15].clear().add_run("Managed €330M annual budget tracking, identifying financial risks and performing variance analysis to formulate strategic recommendations.")
    doc.paragraphs[16].clear()
    doc.paragraphs[17].clear().add_run("Extracted and consolidated financial data using SAP BW/BO, Excel (VBA), and Power Query to optimize monthly budget execution reports.")
    doc.paragraphs[18].clear()
    doc.paragraphs[19].clear().add_run("Automated financial dashboards with Power BI, constructing KPIs that reduced reporting production lead times.")
    
    p20 = doc.paragraphs[20]
    p20.clear()
    p20.add_run("Bridge Asset Management\tAbidjan, Côte d’Ivoire").bold = True
    
    doc.paragraphs[21].clear().add_run("Assistant Portfolio Manager Intern\tMay 2024 - Aug 2024")
    doc.paragraphs[22].clear().add_run("Analyzed corporate financial statements and tracked €137M in investment portfolio performance.")
    
    doc.paragraphs[23].clear().add_run("Analyzed financial statements, cash flows, and balance sheets for major corporate clients to assess performance and financial position.")
    doc.paragraphs[24].clear().add_run("Executed due diligence processes for shareholding transactions, verifying financial documents and compliance.")
    doc.paragraphs[25].clear().add_run("Monitored the daily performance of an investment portfolio exceeding €137M and drafted market analysis summaries for senior management.")
    
    p26 = doc.paragraphs[26]
    p26.clear()
    p26.add_run("Ministry of Europe and Foreign Affairs (MEAE)\tParis, France").bold = True
    
    doc.paragraphs[27].clear().add_run("Budget & Accounting Officer (Contractor)\tJun 2026 - Aug 2026")
    doc.paragraphs[28].clear().add_run("Executed analytical tracking and control of institutional budgets using SAP.")
    
    doc.paragraphs[29].clear().add_run("Ensured strict compliance with public financial regulations while processing budgetary commitments and performing financial reconciliations in SAP.")
    doc.paragraphs[30].clear().add_run("Built dynamic dashboards to optimize the tracking of financial KPIs and accelerate institutional decision-making.")
    
    p32 = doc.paragraphs[32]
    p32.clear()
    p32.add_run("Languages: ").bold = True
    p32.add_run("Fluent in French and English; Conversational Proficiency in Spanish")
    
    p33 = doc.paragraphs[33]
    p33.clear()
    p33.add_run("Technical Skills: ").bold = True
    p33.add_run("Python, VBA, Power Query, Power BI, SAP BW/BO")
    
    p34 = doc.paragraphs[34]
    p34.clear()
    p34.add_run("Certifications & Training: ").bold = True
    p34.add_run("TOEFL ITP: 643/677, Bachibac diploma")
    
    p35 = doc.paragraphs[35]
    p35.clear()
    p35.add_run("Activities: ").bold = True
    p35.add_run("8 years of Capoeira, Played organized Basketball (Prenational level)")
    
    p36 = doc.paragraphs[36]
    p36.clear()
    p36.add_run("Interests: ").bold = True
    p36.add_run("Cinema, Guitar")
    
    for p in doc.paragraphs:
        if p.text == "":
            p._element.getparent().remove(p._element)

    doc.save(output_path)

if __name__ == "__main__":
    template = "/Users/hailesotome/Library/Mobile Documents/com~apple~CloudDocs/CV/Finance Resume Template.docx"
    output = "/Users/hailesotome/Library/Mobile Documents/com~apple~CloudDocs/CV/Tailored_Finance_Resume_PwC.docx"
    create_tailored_resume(template, output)
    print("Resume tailored and saved to:", output)
