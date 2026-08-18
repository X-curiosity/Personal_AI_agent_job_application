import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def set_font(doc, name='Times New Roman', size=11):
    style = doc.styles['Normal']
    font = style.font
    font.name = name
    font.size = Pt(size)

def add_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    add_horizontal_line(p)
    # small space after
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(6)

def add_experience(doc, left_bold, right_normal, left_italic, right_italic):
    # Line 1: Bold left (ALL CAPS), Normal right
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    
    # Add tab stop at 7.5 inches right aligned
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    
    r1 = p.add_run(left_bold.upper())
    r1.bold = True
    p.add_run("\t" + right_normal)
    
    # Line 2: Italic left, Italic right
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    tab_stops2 = p2.paragraph_format.tab_stops
    tab_stops2.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    
    r2_left = p2.add_run(left_italic)
    r2_left.italic = True
    r2_right = p2.add_run("\t" + right_italic)
    r2_right.italic = True

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.25)

def main():
    doc = docx.Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    set_font(doc, 'Times New Roman', 10)
    
    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(0)
    name_run = name_p.add_run("Haile SOTOME".upper())
    name_run.bold = True
    name_run.font.size = Pt(15)
    
    # Contact
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(4)
    contact_p.add_run("Montreuil, France | +33 6 95 59 86 39 | haile.sotome@skema.edu")
    
    # PROFILE
    add_header(doc, "PROFILE")
    profile_p = doc.add_paragraph("Trilingual MSc Corporate Financial Management student with a strong quantitative foundation and hands-on experience in financial analysis and portfolio management. Highly rigorous and curious, with advanced analytical skills (VBA, Power Query, SAP). Seeking a challenging internship in Corporate Finance / M&A to leverage my financial expertise in supporting market analysis, company valuations, and deal execution.")
    profile_p.paragraph_format.space_after = Pt(2)
    
    # EDUCATION
    add_header(doc, "EDUCATION")
    
    add_experience(doc, 
                   "SKEMA Business School - Soochow University", "Suzhou, China / Paris, France",
                   "Masters PGE: Specialization in Corporate Financial Management (CFM)", "Expected 2027")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Relevant Coursework: ").bold = True
    p.add_run("Corporate Valuation, Financial Modelling, Cash Flow Forecasting, Risk Management, Portfolio Management")
    
    add_experience(doc, 
                   "SKEMA Business School - North Carolina State University (NCSU)", "Lille, France / Raleigh, US",
                   "Masters PGE: Specialization in Finance & Quants", "2023 - 2026")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Relevant Coursework: ").bold = True
    p.add_run("Maths for Finance, Big Data Analysis, Derivatives, Structured Products, Python/VBA")

    add_experience(doc,
                   "Lycée Kléber - CPGE", "Strasbourg, France",
                   "Intensive Economics Preparatory Classes for Grandes Écoles (CPGE)", "2021 - 2023")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run("Coursework: ").bold = True
    p.add_run("Applied Mathematics, Economics, Geopolitics, Philosophy, English, Spanish")
                   
    # WORK EXPERIENCE
    add_header(doc, "WORK & LEADERSHIP EXPERIENCE")
    
    add_experience(doc,
                   "Université Paris-Est Créteil (UPEC)", "Créteil, France",
                   "Financial Analyst Intern, Financial Affairs Division", "May 2025 - December 2025")
    add_bullet(doc, "Managed €330M annual budget tracking, identifying financial risks and performing variance analysis to formulate strategic recommendations.")
    add_bullet(doc, "Extracted and consolidated financial data using SAP BW/BO, Excel (VBA), and Power Query to optimize monthly budget execution reports.")
    add_bullet(doc, "Automated financial dashboards with Power BI, constructing KPIs that reduced reporting production lead times.")
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    
    add_experience(doc,
                   "Bridge Asset Management", "Abidjan, Côte d’Ivoire",
                   "Assistant Portfolio Manager Intern", "May 2024 - August 2024")
    add_bullet(doc, "Analyzed financial statements, cash flows, and balance sheets for major corporate clients to assess performance and financial position.")
    add_bullet(doc, "Executed due diligence processes for shareholding transactions, verifying financial documents and compliance.")
    add_bullet(doc, "Monitored the daily performance of an investment portfolio exceeding €137M and drafted market analysis summaries for senior management.")
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    add_experience(doc,
                   "Ministry of Europe and Foreign Affairs (MEAE)", "Paris, France",
                   "Budget & Accounting Officer (Contractor)", "June 2026 - August 2026")
    add_bullet(doc, "Ensured strict compliance with public financial regulations while processing budgetary commitments and performing financial reconciliations in SAP.")
    add_bullet(doc, "Built dynamic dashboards to optimize the tracking of financial KPIs and accelerate institutional decision-making.")
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # SKILLS & INTERESTS
    add_header(doc, "SKILLS, ACTIVITIES & INTERESTS")
    
    def add_skill_line(doc, label, content):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run(label + ": ").bold = True
        p.add_run(content)
        
    add_skill_line(doc, "Languages", "Fluent in French and English; Conversational Proficiency in Spanish")
    add_skill_line(doc, "Technical Skills", "Python, VBA, Power Query, Power BI, SAP BW/BO")
    add_skill_line(doc, "Certifications & Training", "TOEFL ITP: 643/677, Bachibac diploma")
    add_skill_line(doc, "Activities", "8 years of Capoeira, Played organized Basketball (Prenational level)")
    add_skill_line(doc, "Interests", "Cinema, Guitar")
    
    doc.save("/Users/hailesotome/Library/Mobile Documents/com~apple~CloudDocs/CV/Tailored_Finance_Resume_PwC_Final.docx")

if __name__ == "__main__":
    main()
    print("Successfully saved to Tailored_Finance_Resume_PwC_Final.docx")
