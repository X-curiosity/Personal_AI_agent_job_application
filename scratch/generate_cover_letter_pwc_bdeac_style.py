import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(doc, font_name, size):
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = size
    font.color.rgb = RGBColor(0, 0, 0)

def main():
    doc = docx.Document()
    
    # Set standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    set_font(doc, 'Times New Roman', 11)
    
    # ----------------------------------------------------
    # CANDIDATE INFO (Top Left)
    # ----------------------------------------------------
    cand_p = doc.add_paragraph()
    cand_p.paragraph_format.space_after = Pt(0)
    cand_p.add_run("Haile SOTOME\n").bold = True
    cand_p.add_run("Montreuil, France\n")
    cand_p.add_run("+33 6 95 59 86 39\n")
    cand_p.add_run("haile.sotome@skema.edu")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # EMPLOYER INFO (Left)
    # ----------------------------------------------------
    emp_p = doc.add_paragraph()
    emp_p.paragraph_format.space_after = Pt(0)
    emp_p.add_run("À l'attention de l'Équipe Recrutement\n")
    emp_p.add_run("PwC France\n")
    emp_p.add_run("63 Rue de Villiers\n")
    emp_p.add_run("92200 Neuilly-sur-Seine, France")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # DATE (Right Aligned)
    # ----------------------------------------------------
    date_p = doc.add_paragraph("Fait à Montreuil, le 27 juillet 2026")
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.paragraph_format.space_after = Pt(24)
    
    # ----------------------------------------------------
    # SUBJECT
    # ----------------------------------------------------
    subj_p = doc.add_paragraph()
    subj_p.add_run("Objet : Candidature pour le Stage Corporate Finance (Janvier 2027) F/H").bold = True
    subj_p.paragraph_format.space_after = Pt(24)
    
    # ----------------------------------------------------
    # GREETING
    # ----------------------------------------------------
    greet_p = doc.add_paragraph("Madame, Monsieur,")
    greet_p.paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # BODY PARAGRAPHS (Based on BDEAC exactly)
    # ----------------------------------------------------
    p1 = doc.add_paragraph("PwC joue un rôle d'intégration et d'accélérateur de mouvement majeur à l'échelle mondiale et en France. Face aux dynamiques de transformation des entreprises et d'évaluation des risques financiers propres à un cabinet pluridisciplinaire d'excellence, je souhaite mettre ma double compétence quantitative et analytique au service de vos équipes.")
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.space_after = Pt(12)

    p2 = doc.add_paragraph("Actuellement en Master 2 au sein du MSc Corporate Financial Management (CFM) à SKEMA Business School, sur le campus de Suzhou en Chine en partenariat avec la Soochow University, j'approfondis mon expertise en finance d'entreprise après un Master 1 Finance & Quants (Financial Markets & Instruments) effectué en partenariat avec North Carolina State University (NCSU) à Raleigh. Ce parcours d'excellence, entamé après deux années de Classes Préparatoires (ECG) au Lycée Kléber, m'a permis de maîtriser des concepts avancés tels que la valorisation d'entreprise, le private equity, la gestion des risques financiers et la finance internationale.")
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.space_after = Pt(12)
    
    p3 = doc.add_paragraph("Au-delà de cette formation théorique rigoureuse, je possède une expérience opérationnelle déjà ancrée dans la finance et l'analyse stratégique :")
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.paragraph_format.space_after = Pt(6)

    # Bullets
    b1 = doc.add_paragraph(style='List Bullet')
    b1.add_run("En gestion d'actifs (Abidjan, Côte d'Ivoire) : ").bold = True
    b1.add_run("En tant qu'assistant gestionnaire de portefeuille chez Bridge Asset Management, j'ai analysé les comptes sociaux et consolidés de grands comptes, assuré le suivi de performance de portefeuilles d'investissements et participé activement aux processus de due diligences.")
    b1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    b1.paragraph_format.space_after = Pt(6)

    b2 = doc.add_paragraph(style='List Bullet')
    b2.add_run("En analyse financière et gestion publique (Paris) : ").bold = True
    b2.add_run("Actuellement contrôleur de gestion en CDD au Ministère de l'Europe et des Affaires Étrangères (MEAE), je travaille sur la modélisation des trajectoires budgétaires pluriannuelles stratégiques. Auparavant, à la Direction des Affaires Financières de l'UPEC, j'ai participé au suivi budgétaire d'un établissement de 330 M€, géré des clôtures comptables et automatisé des processus de reporting.")
    b2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    b2.paragraph_format.space_after = Pt(12)

    p4 = doc.add_paragraph("Ces expériences m'ont permis de développer une expertise technique pointue dans l'exploitation de données complexes et la conception d'outils d'aide à la décision (SAP BW/BO, Power Query, Power BI, VBA et Python). Parfaitement trilingue (natif en français et en anglais, courant en espagnol grâce au Bachibac), je dispose de l'agilité multiculturelle indispensable pour évoluer au sein d'un réseau mondial tel que PwC.")
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4.paragraph_format.space_after = Pt(12)
    
    p5 = doc.add_paragraph("Intégrer PwC représente l'opportunité de mettre ma rigueur analytique au profit de missions à fort impact. Dans le cadre de mon stage de fin d'études de mon MSc CFM, je suis disponible pour une durée de 6 mois à compter de janvier 2027.")
    p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p5.paragraph_format.space_after = Pt(12)

    # ----------------------------------------------------
    # CLOSING PARAGRAPH
    # ----------------------------------------------------
    p6 = doc.add_paragraph("Je serais ravi de vous exposer de vive voix la valeur ajoutée que je peux apporter à vos équipes lors d'un entretien.")
    p6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p6.paragraph_format.space_after = Pt(12)

    # ----------------------------------------------------
    # SIGN-OFF
    # ----------------------------------------------------
    sign_p = doc.add_paragraph("Je vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées.")
    sign_p.paragraph_format.space_after = Pt(24)
    
    name_sign_p = doc.add_paragraph("Haile SOTOME")
    name_sign_p.paragraph_format.space_after = Pt(0)
    
    # Save document
    doc.save("/Users/hailesotome/Library/Mobile Documents/com~apple~CloudDocs/CV/Cover_Letter_PwC_BDEAC_Style.docx")

if __name__ == "__main__":
    main()
    print("Cover letter successfully saved to Cover_Letter_PwC_BDEAC_Style.docx")
