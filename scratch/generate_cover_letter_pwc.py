import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def set_font(doc, font_name, size):
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = size
    # Explicitly set color to black to avoid 'all white' bug
    font.color.rgb = RGBColor(0, 0, 0)

def main():
    doc = docx.Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    set_font(doc, 'Times New Roman', 11)
    
    # ----------------------------------------------------
    # MASTHEAD (Matches Resume)
    # ----------------------------------------------------
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(0)
    name_run = name_p.add_run("Haile SOTOME".upper())
    name_run.bold = True
    name_run.font.size = Pt(15)
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(24)
    contact_run = contact_p.add_run("Montreuil, France | +33 6 95 59 86 39 | haile.sotome@skema.edu")
    contact_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # ----------------------------------------------------
    # DATE
    # ----------------------------------------------------
    date_p = doc.add_paragraph("Le 27 juillet 2026")
    date_p.paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # EMPLOYER INFO
    # ----------------------------------------------------
    emp_p = doc.add_paragraph()
    emp_p.paragraph_format.space_after = Pt(12)
    emp_p.add_run("Équipe Recrutement PwC France\n").font.color.rgb = RGBColor(0, 0, 0)
    emp_p.add_run("PwC France\n").font.color.rgb = RGBColor(0, 0, 0)
    emp_p.add_run("63 Rue de Villiers\n").font.color.rgb = RGBColor(0, 0, 0)
    emp_p.add_run("92200 Neuilly-sur-Seine, France").font.color.rgb = RGBColor(0, 0, 0)
    
    # ----------------------------------------------------
    # SUBJECT
    # ----------------------------------------------------
    subj_p = doc.add_paragraph()
    subj_p.paragraph_format.space_after = Pt(12)
    subj_run = subj_p.add_run("Objet : Candidature pour le Stage Corporate Finance (Janvier 2027) F/H")
    subj_run.bold = True
    subj_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # ----------------------------------------------------
    # GREETING
    # ----------------------------------------------------
    greet_p = doc.add_paragraph("Madame, Monsieur,")
    greet_p.paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # BODY PARAGRAPHS
    # ----------------------------------------------------
    p1 = doc.add_paragraph()
    p1.add_run("Actuellement en Master 2 au sein du MSc Corporate Financial Management à SKEMA Business School, je souhaite mettre ma double compétence quantitative et analytique au service de vos équipes en tant que stagiaire Corporate Finance à partir de janvier 2027. Animé par l'ambition de PwC d'être un véritable \"accélérateur de mouvement\" et me reconnaissant pleinement dans la promesse \"Grow here. Go further.\", je suis convaincu que ma curiosité, ma rigueur et ma passion pour la finance d'entreprise me permettront d'incarner les comportements du PwC Professional.").font.color.rgb = RGBColor(0, 0, 0)
    p1.paragraph_format.space_after = Pt(12)
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p2 = doc.add_paragraph()
    p2.add_run("Au-delà de ma formation théorique d'excellence, qui m'a permis de maîtriser les concepts avancés tels que la valorisation d'entreprise et la gestion des risques financiers, je possède une expérience opérationnelle directement transposable aux exigences de PwC. En tant qu'assistant gestionnaire de portefeuille chez Bridge Asset Management, j'ai participé activement aux processus de due diligence pour des opérations capitalistiques et analysé les comptes sociaux et consolidés de grands comptes afin de soutenir le suivi de performance d'un portefeuille de 137 M€. De plus, mon expérience en tant que contrôleur de gestion / analyste financier à l'UPEC m'a amené à automatiser des reportings complexes et à construire des tableaux de bord dynamiques via Power BI, Excel VBA et SAP BW/BO, permettant d'optimiser la prise de décision stratégique sur un budget de 330 M€.").font.color.rgb = RGBColor(0, 0, 0)
    p2.paragraph_format.space_after = Pt(12)
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p3 = doc.add_paragraph()
    p3.add_run("Parfaitement trilingue (français, anglais, espagnol) et fort d'un parcours académique international (campus de Suzhou en Chine et NCSU aux États-Unis), je dispose de l'agilité multiculturelle et de la curiosité intellectuelle indispensables pour évoluer efficacement au sein d'un cabinet pluridisciplinaire d'envergure mondiale tel que PwC. Ces expériences m'ont permis de développer une expertise technique pointue dans l'exploitation de données complexes, que je suis prêt à mobiliser pour accompagner au mieux vos clients.").font.color.rgb = RGBColor(0, 0, 0)
    p3.paragraph_format.space_after = Pt(12)
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ----------------------------------------------------
    # CLOSING PARAGRAPH
    # ----------------------------------------------------
    p4 = doc.add_paragraph()
    p4.add_run("Intégrer PwC représente l'opportunité de mettre ma rigueur analytique et mon énergie au profit de missions à fort impact. Je serais ravi de vous exposer de vive voix la valeur ajoutée que je peux apporter à vos équipes lors d'un entretien.").font.color.rgb = RGBColor(0, 0, 0)
    p4.paragraph_format.space_after = Pt(24)
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ----------------------------------------------------
    # SIGN-OFF
    # ----------------------------------------------------
    sign_p = doc.add_paragraph()
    sign_p.add_run("Je vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées.").font.color.rgb = RGBColor(0, 0, 0)
    sign_p.paragraph_format.space_after = Pt(24)
    
    name_sign_p = doc.add_paragraph()
    name_sign_p.add_run("Haile SOTOME").font.color.rgb = RGBColor(0, 0, 0)
    name_sign_p.paragraph_format.space_after = Pt(0)
    
    # Save document
    doc.save("/Users/hailesotome/Library/Mobile Documents/com~apple~CloudDocs/CV/Cover_Letter_PwC.docx")

if __name__ == "__main__":
    main()
    print("Cover letter successfully saved to Cover_Letter_PwC.docx")
