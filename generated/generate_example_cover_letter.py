import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font(doc, font_name, size):
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = size
    font.color.rgb = RGBColor(0, 0, 0)

def generate_cover_letter(output_path: str):
    doc = docx.Document()
    
    # ----------------------------------------------------
    # PAGE MARGINS: 1 inch on all sides
    # ----------------------------------------------------
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    set_font(doc, 'Times New Roman', Pt(11))
    
    # ----------------------------------------------------
    # CANDIDATE INFO (Top-Left, Name Bold, Rest Plain)
    # ----------------------------------------------------
    cand_p = doc.add_paragraph()
    cand_p.paragraph_format.space_after = Pt(6)
    cand_name = cand_p.add_run("Haile SOTOME\n")
    cand_name.bold = True
    cand_p.add_run("Montreuil, France\n")
    cand_p.add_run("+33 6 95 59 86 39\n")
    cand_p.add_run("haile.sotome@skema.edu")
    
    # ----------------------------------------------------
    # DATE (Left-Aligned)
    # ----------------------------------------------------
    date_p = doc.add_paragraph("August 18, 2026")
    date_p.paragraph_format.space_after = Pt(6)
    
    # ----------------------------------------------------
    # EMPLOYER INFO (Left-Aligned)
    # ----------------------------------------------------
    emp_p = doc.add_paragraph()
    emp_p.paragraph_format.space_after = Pt(6)
    emp_p.add_run("Hiring Manager\n")
    emp_p.add_run("UBS AG\n")
    emp_p.add_run("Bahnhofstrasse 45\n")
    emp_p.add_run("8001 Zurich, Switzerland")
    
    # ----------------------------------------------------
    # SUBJECT LINE (Bold)
    # ----------------------------------------------------
    subj_p = doc.add_paragraph()
    subj_p.paragraph_format.space_after = Pt(6)
    subj_run = subj_p.add_run("RE: Junior Financial Analyst Position")
    subj_run.bold = True
    
    # ----------------------------------------------------
    # GREETING
    # ----------------------------------------------------
    greet_p = doc.add_paragraph("Dear Hiring Manager,")
    greet_p.paragraph_format.space_after = Pt(6)
    
    # ----------------------------------------------------
    # OPENING PARAGRAPH (Hook)
    # ----------------------------------------------------
    p1 = doc.add_paragraph(
        "As a trilingual MSc Corporate Financial Management student with hands-on experience managing a "
        "€330M public budget and monitoring a €137M investment portfolio, I am writing to express my strong "
        "interest in the Junior Financial Analyst position at UBS. UBS's commitment to sustainable finance "
        "and leadership in global wealth management align directly with my ambition to apply rigorous "
        "quantitative analysis within a world-class institution."
    )
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.space_after = Pt(6)
    
    # ----------------------------------------------------
    # BODY PARAGRAPH 1 (Financial Analysis & Budget Management)
    # ----------------------------------------------------
    p2 = doc.add_paragraph(
        "In my recent role at the Ministry of Europe and Foreign Affairs and my prior internship at "
        "Université Paris-Est Créteil, I honed my ability to manage complex public finances. I ensured compliance "
        "and performed financial reconciliations in SAP, while also tracking a €330M annual budget and formulating "
        "strategic recommendations based on variance analysis. By automating dashboards with Power BI and consolidating "
        "datasets using VBA and Power Query, I reduced reporting lead times and optimized the tracking of financial KPIs. "
        "This data-driven approach to analysis and process improvement underpins the advisory functions at UBS."
    )
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.space_after = Pt(6)
    
    # ----------------------------------------------------
    # BODY PARAGRAPH 2 (Portfolio Management & Due Diligence)
    # ----------------------------------------------------
    p3 = doc.add_paragraph(
        "At Bridge Asset Management in Abidjan, I analyzed financial statements, cash flows, and balance sheets "
        "for major corporate clients while monitoring the daily performance of a €137M investment portfolio. "
        "I executed due diligence for shareholding transactions and drafted market analysis summaries for senior "
        "management, deepening my understanding of asset valuation and compliance frameworks directly transferable "
        "to UBS's analyst functions."
    )
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.paragraph_format.space_after = Pt(6)
    
    # ----------------------------------------------------
    # CLOSING PARAGRAPH
    # ----------------------------------------------------
    p4 = doc.add_paragraph(
        "My rigorous quantitative training from intensive Classes Préparatoires and my Finance and Quants "
        "master's degree, combined with multilingual fluency in French, English, and Spanish, and proven experience "
        "in financial analysis would allow me to contribute meaningfully to UBS's team from day one. I would "
        "welcome the opportunity to discuss how my background aligns with your needs. Thank you for your consideration."
    )
    p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4.paragraph_format.space_after = Pt(6)
    
    # ----------------------------------------------------
    # SIGN-OFF (24pt space before name)
    # ----------------------------------------------------
    sign_p = doc.add_paragraph("Sincerely,")
    sign_p.paragraph_format.space_after = Pt(12)
    
    name_p = doc.add_paragraph("Haile SOTOME")
    name_p.paragraph_format.space_after = Pt(0)
    
    # Ensure directory exists and save document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Cover letter successfully saved to: {output_path}")

def main():
    output_path = "/Users/hailesotome/.gemini/antigravity/scratch/Personal_AI_agent_job_application/generated/Cover_Letter_UBS_Junior_Financial_Analyst.docx"
    generate_cover_letter(output_path)

if __name__ == "__main__":
    main()
